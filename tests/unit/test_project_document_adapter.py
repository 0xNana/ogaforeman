from __future__ import annotations

from io import BytesIO

import pytest
import xlwt
from docx import Document
from openpyxl import Workbook
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from app.domain.project_import import SourceType
from app.services.project_source_adapter import (
    ProjectDocumentAdapter,
    StructuredTextInputError,
)


def _docx_bytes() -> bytes:
    document = Document()
    document.add_heading("Ridge House", level=1)
    document.add_paragraph("Task: Foundation")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Material"
    table.cell(0, 1).text = "Quantity"
    table.cell(1, 0).text = "Cement"
    table.cell(1, 1).text = "100 bags"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _xlsx_bytes() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Programme"
    sheet.append(["Task", "Due"])
    sheet.append(["Foundation", "2026-09-10"])
    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def _xls_bytes() -> bytes:
    workbook = xlwt.Workbook()
    sheet = workbook.add_sheet("Programme")
    for column, value in enumerate(("Task", "Due")):
        sheet.write(0, column, value)
    for column, value in enumerate(("Foundation", "2026-09-10")):
        sheet.write(1, column, value)
    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def _pdf_bytes() -> bytes:
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): writer._add_object(font)})}
    )
    content = DecodedStreamObject()
    content.set_data(b"BT /F1 12 Tf 72 720 Td (Task: Foundation) Tj ET")
    page[NameObject("/Contents")] = writer._add_object(content)
    output = BytesIO()
    writer.write(output)
    return output.getvalue()


@pytest.mark.parametrize(
    ("name", "content", "source_type", "expected"),
    [
        ("ridge-plan.docx", _docx_bytes(), SourceType.FILE, "Cement"),
        ("ridge-plan.pdf", _pdf_bytes(), SourceType.FILE, "Foundation"),
        ("ridge-plan.xlsx", _xlsx_bytes(), SourceType.SPREADSHEET, "Programme"),
        ("ridge-plan.xls", _xls_bytes(), SourceType.SPREADSHEET, "Foundation"),
        (
            "ridge-plan.csv",
            b"Task,Due\r\nFoundation,2026-09-10\r\n",
            SourceType.SPREADSHEET,
            "Foundation",
        ),
    ],
)
def test_document_adapter_extracts_supported_project_files(
    name: str,
    content: bytes,
    source_type: SourceType,
    expected: str,
) -> None:
    source = ProjectDocumentAdapter(name=name, source_type=source_type).load(content)

    assert source.name == name
    assert source.source_type is source_type
    assert "Foundation" in source.text
    assert expected in source.text


def test_document_adapter_rejects_mismatched_empty_and_oversized_files() -> None:
    with pytest.raises(StructuredTextInputError):
        ProjectDocumentAdapter(name="ridge-plan.xlsx", source_type=SourceType.FILE).load(
            _xlsx_bytes()
        )
    with pytest.raises(StructuredTextInputError):
        ProjectDocumentAdapter(name="ridge-plan.pdf", source_type=SourceType.FILE).load(b"")
    with pytest.raises(StructuredTextInputError):
        ProjectDocumentAdapter(name="ridge-plan.csv", source_type=SourceType.SPREADSHEET).load(
            b"x" * (ProjectDocumentAdapter.MAX_FILE_BYTES + 1)
        )


def test_document_adapter_rejects_a_pdf_without_extractable_text() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    output = BytesIO()
    writer.write(output)

    with pytest.raises(StructuredTextInputError, match="extractable text"):
        ProjectDocumentAdapter(name="scan.pdf", source_type=SourceType.FILE).load(output.getvalue())


def test_document_adapter_rejects_binary_content_renamed_as_csv() -> None:
    with pytest.raises(StructuredTextInputError, match="binary content"):
        ProjectDocumentAdapter(
            name="ridge-plan.csv",
            source_type=SourceType.SPREADSHEET,
        ).load(b"Task,Due\nFoundation,\x00\x01\x02")
